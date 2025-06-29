#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import json
import argparse
import logging
import time
import requests
import base64
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/sentinel1_reporting_agent.log"),
        logging.StreamHandler()
    ]
)

class Sentinel1ReportingAgent:
    """
    Agent for generating summary reports from Sentinel-1 model training results.
    """
    
    def __init__(self, server_url, token):
        """
        Initialize the Sentinel1ReportingAgent.
        
        Args:
            server_url (str): URL of the MCP server
            token (str): Authentication token for the MCP server
        """
        self.server_url = server_url
        self.token = token
        self.agent_id = "sentinel1-reporting-agent"
        self.logger = logging.getLogger("sentinel1_reporting_agent")
        
        # Validate token
        self._validate_token()
        self.logger.info(f"Sentinel-1 Reporting Agent initialized with ID: {self.agent_id}")
    
    def _validate_token(self):
        """Validate the agent token with the MCP server."""
        try:
            response = requests.post(
                f"{self.server_url}/api/validate-token",
                json={
                    "agent_id": self.agent_id,
                    "agent_type": "sentinel1-reporting",
                    "token": self.token
                }
            )
            response.raise_for_status()
            self.logger.info("Agent token validated successfully")
        except requests.exceptions.RequestException as e:
            self.logger.error(f"Failed to validate token: {e}")
            raise
    
    def get_tasks(self):
        """
        Get tasks assigned to this agent from the MCP server.
        
        Returns:
            list: List of tasks assigned to this agent
        """
        try:
            response = requests.get(
                f"{self.server_url}/api/tasks/{self.agent_id}",
                params={"token": self.token}
            )
            response.raise_for_status()
            result = response.json()
            self.logger.info(f"Retrieved {len(result.get('tasks', []))} tasks")
            return result.get('tasks', [])
        except requests.exceptions.RequestException as e:
            self.logger.error(f"Failed to get tasks: {e}")
            return []
    
    def update_task_status(self, task_id, status, max_retries=4):
        """
        Update the status of a task on the MCP server.
        
        Args:
            task_id (str): ID of the task to update
            status (str): New status of the task
            max_retries (int): Maximum number of retry attempts
        
        Returns:
            bool: True if the update was successful, False otherwise
        """
        for attempt in range(1, max_retries + 1):
            try:
                self.logger.info(f"Updating task {task_id} status to '{status}' (attempt {attempt}/{max_retries})")
                response = requests.put(
                    f"{self.server_url}/api/tasks/{task_id}/status",
                    params={"token": self.token},
                    json={"status": status}
                )
                response.raise_for_status()
                return True
            except requests.exceptions.RequestException as e:
                self.logger.warning(f"Failed to update task status: {e}")
                if attempt < max_retries:
                    time.sleep(2 ** (attempt - 1))  # Exponential backoff
                else:
                    self.logger.error(f"Failed to update task {task_id} status after {max_retries} attempts")
                    return False
    
    def process_task(self, task):
        """
        Process a reporting task.
        
        Args:
            task (dict): Task information
        """
        task_id = task["id"]
        self.logger.info(f"Processing task {task_id}: {task.get('title', 'No title')}")
        
        # Update task status to in_progress
        if not self.update_task_status(task_id, "in_progress"):
            return
        
        try:
            # Extract metadata from task
            metadata = task.get("metadata", "{}")
            
            # Parse metadata if it's a string
            if isinstance(metadata, str):
                try:
                    metadata = json.loads(metadata)
                except json.JSONDecodeError:
                    self.logger.error(f"Failed to parse metadata as JSON: {metadata}")
                    metadata = {}
            
            model_dir = metadata.get("model_dir", "")
            output_dir = metadata.get("output_dir", "")
            
            if not model_dir or not os.path.exists(model_dir):
                self.logger.error(f"Invalid model directory: {model_dir}")
                self.update_task_status(task_id, "failed")
                return
            
            if not output_dir:
                output_dir = os.path.join(os.path.dirname(model_dir), "reports")
            
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate report
            report_path = self.generate_report(model_dir, output_dir)
            
            if report_path:
                self.logger.info(f"Report generated successfully: {report_path}")
                self.update_task_status(task_id, "completed")
            else:
                self.logger.error("Failed to generate report")
                self.update_task_status(task_id, "failed")
        
        except Exception as e:
            self.logger.error(f"Error processing task: {e}", exc_info=True)
            self.update_task_status(task_id, "failed")
    
    def generate_report(self, model_dir, output_dir):
        """
        Generate a summary report from model training results.
        
        Args:
            model_dir (str): Directory containing model results
            output_dir (str): Directory to save the report
        
        Returns:
            str: Path to the generated report, or None if generation failed
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Add title
            doc.add_heading('Sentinel-1 Model Training Results', 0)
            
            # Add timestamp
            doc.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # Add overall metrics section
            doc.add_heading('Overall Performance Metrics', level=1)
            
            # Read and parse overall.json
            overall_path = os.path.join(model_dir, 'overall.json')
            if os.path.exists(overall_path):
                with open(overall_path, 'r') as f:
                    overall_metrics = json.load(f)
                
                # Add metrics table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Metric'
                header_cells[1].text = 'Value'
                
                # Add metrics rows
                for metric, value in overall_metrics.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = metric
                    row_cells[1].text = f"{value:.4f}" if isinstance(value, (int, float)) else str(value)
                
                # Add explanation
                doc.add_paragraph('')
                doc.add_heading('Metrics Explanation', level=2)
                doc.add_paragraph('• Accuracy: The proportion of correctly classified samples out of all samples.')
                doc.add_paragraph('• IoU (Intersection over Union): Measures the overlap between predicted and ground truth segmentations.')
                doc.add_paragraph('• Precision: The ability of the model to identify only relevant instances.')
                doc.add_paragraph('• Recall: The ability of the model to find all relevant instances.')
                doc.add_paragraph('• F1-score: The harmonic mean of precision and recall.')
                doc.add_paragraph('• Micro metrics: Calculated by aggregating the contributions of all classes.')
                doc.add_paragraph('• Macro metrics: Calculated by taking the average of the metrics computed for each class.')
            else:
                doc.add_paragraph('Overall metrics file (overall.json) not found.')
            
            # Add accuracy graph
            doc.add_heading('Training and Validation Accuracy', level=1)
            accuracy_graph_path = os.path.join(model_dir, 'accuracy_graph.png')
            if os.path.exists(accuracy_graph_path):
                doc.add_picture(accuracy_graph_path, width=Inches(6))
                doc.add_paragraph('The accuracy graph shows the model\'s classification accuracy on both training and validation datasets over the course of training epochs. Higher values indicate better performance, and the convergence of training and validation curves suggests good generalization.')
            else:
                doc.add_paragraph('Accuracy graph file (accuracy_graph.png) not found.')
            
            # Add loss graph
            doc.add_heading('Training and Validation Loss', level=1)
            loss_graph_path = os.path.join(model_dir, 'loss_graph.png')
            if os.path.exists(loss_graph_path):
                doc.add_picture(loss_graph_path, width=Inches(6))
                doc.add_paragraph('The loss graph shows the model\'s loss function value on both training and validation datasets over the course of training epochs. Lower values indicate better performance, and the convergence of training and validation curves suggests good generalization without overfitting.')
            else:
                doc.add_paragraph('Loss graph file (loss_graph.png) not found.')
            
            # Add confusion matrix
            doc.add_heading('Confusion Matrix', level=1)
            conf_mat_path = os.path.join(model_dir, 'conf_mat_picture.png')
            if os.path.exists(conf_mat_path):
                doc.add_picture(conf_mat_path, width=Inches(6))
                doc.add_paragraph('The confusion matrix shows the distribution of predicted classes versus actual classes. The diagonal elements represent correctly classified instances, while off-diagonal elements represent misclassifications. Higher values along the diagonal and lower values elsewhere indicate better model performance.')
            else:
                doc.add_paragraph('Confusion matrix file (conf_mat_picture.png) not found.')
            
            # Add per-class metrics if available
            per_class_path = os.path.join(model_dir, 'per_class.json')
            if os.path.exists(per_class_path):
                doc.add_heading('Per-Class Performance', level=1)
                with open(per_class_path, 'r') as f:
                    per_class_metrics = json.load(f)
                
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Class'
                header_cells[1].text = 'IoU'
                header_cells[2].text = 'Precision'
                header_cells[3].text = 'Recall'
                header_cells[4].text = 'F1-score'
                
                # Add class metrics rows
                for class_id, metrics in per_class_metrics.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = class_id
                    row_cells[1].text = f"{metrics.get('IoU', 'N/A'):.4f}" if isinstance(metrics.get('IoU'), (int, float)) else str(metrics.get('IoU', 'N/A'))
                    row_cells[2].text = f"{metrics.get('Precision', 'N/A'):.4f}" if isinstance(metrics.get('Precision'), (int, float)) else str(metrics.get('Precision', 'N/A'))
                    row_cells[3].text = f"{metrics.get('Recall', 'N/A'):.4f}" if isinstance(metrics.get('Recall'), (int, float)) else str(metrics.get('Recall', 'N/A'))
                    row_cells[4].text = f"{metrics.get('F1-score', 'N/A'):.4f}" if isinstance(metrics.get('F1-score'), (int, float)) else str(metrics.get('F1-score', 'N/A'))
            
            # Add conclusion
            doc.add_heading('Conclusion', level=1)
            
            # Generate conclusion based on overall metrics
            if os.path.exists(overall_path):
                with open(overall_path, 'r') as f:
                    overall_metrics = json.load(f)
                
                accuracy = overall_metrics.get('Accuracy', 0)
                f1_score = overall_metrics.get('MACRO_F1-score', 0)
                
                if accuracy > 0.8:
                    performance = "excellent"
                elif accuracy > 0.7:
                    performance = "good"
                elif accuracy > 0.6:
                    performance = "moderate"
                else:
                    performance = "needs improvement"
                
                doc.add_paragraph(f"The Sentinel-1 model demonstrates {performance} performance with an overall accuracy of {accuracy:.2%} and a macro F1-score of {f1_score:.2%}. ")
                
                if performance in ["excellent", "good"]:
                    doc.add_paragraph("The model shows strong generalization capabilities and is suitable for deployment in production environments.")
                elif performance == "moderate":
                    doc.add_paragraph("The model shows reasonable performance but may benefit from additional training data or hyperparameter tuning to improve accuracy.")
                else:
                    doc.add_paragraph("The model's performance indicates that further optimization is needed. Consider revisiting the training data quality, model architecture, or hyperparameters.")
            
            # Save the document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"sentinel1_model_report_{timestamp}.docx"
            report_path = os.path.join(output_dir, report_filename)
            doc.save(report_path)
            
            return report_path
        
        except Exception as e:
            self.logger.error(f"Error generating report: {e}", exc_info=True)
            return None
    
    def run(self):
        """Run the agent to process tasks."""
        self.logger.info("Starting Sentinel1ReportingAgent")
        
        while True:
            try:
                # Get tasks
                tasks = self.get_tasks()
                self.logger.info(f"Retrieved {len(tasks)} tasks")
                
                # Process tasks
                for task in tasks:
                    self.process_task(task)
                
                # Sleep before checking for new tasks
                time.sleep(10)
            
            except Exception as e:
                self.logger.error(f"Error in agent loop: {e}", exc_info=True)
                time.sleep(30)  # Sleep longer on error

def main():
    """Main function to run the Sentinel1ReportingAgent."""
    parser = argparse.ArgumentParser(description='Sentinel-1 Reporting Agent')
    parser.add_argument('--token', required=True, help='Authentication token for the MCP server')
    parser.add_argument('--server-url', default='http://localhost:8080', help='URL of the MCP server')
    
    args = parser.parse_args()
    
    agent = Sentinel1ReportingAgent(args.server_url, args.token)
    agent.run()

if __name__ == '__main__':
    main()
